from docx import Document
import os
import pypandoc
import tempfile
from io import BytesIO

def paragraph_to_markdown(paragraph) -> str:
    """Chuyển đổi một đoạn văn Docx sang Markdown bằng Pandoc."""
    if not paragraph.text.strip():
        return ""
        
    # Tạo docx tạm chứa 1 đoạn văn
    doc = Document()
    new_p = doc.add_paragraph()
    for run in paragraph.runs:
        r = new_p.add_run(run.text)
        r.bold = run.bold
        r.italic = run.italic
        r.underline = run.underline
        
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name
        
    try:
        md = pypandoc.convert_file(tmp_path, 'md', format='docx')
        return md.strip()
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

def markdown_to_paragraph(md_text: str, target_paragraph):
    """Chuyển đổi Markdown ngược về Docx và cập nhật vào paragraph."""
    if not md_text.strip():
        target_paragraph.text = ""
        return

    with tempfile.NamedTemporaryFile(suffix='.md', mode='w', encoding='utf-8', delete=False) as tmp:
        tmp.write(md_text)
        tmp_path = tmp.name
        
    docx_tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    docx_tmp.close()
    docx_tmp_path = docx_tmp.name
    
    try:
        pypandoc.convert_file(tmp_path, 'docx', format='md', outputfile=docx_tmp_path)
        
        # Load lại docx tạm và lấy các runs
        new_doc = Document(docx_tmp_path)
        if new_doc.paragraphs:
            # Xóa các runs cũ
            target_paragraph.text = ""
            # Sao chép các runs mới
            for run in new_doc.paragraphs[0].runs:
                r = target_paragraph.add_run(run.text)
                r.bold = run.bold
                r.italic = run.italic
                r.underline = run.underline
    finally:
        if os.path.exists(tmp_path): os.remove(tmp_path)
        if os.path.exists(docx_tmp_path): os.remove(docx_tmp_path)

def iter_paragraphs(parent):
    """
    Duyệt đệ quy tất cả các đoạn văn trong Document hoặc Table Cell.
    Giúp trích xuất text từ cả Paragraphs và Tables.
    """
    from docx.document import Document as _Document
    from docx.table import _Cell, Table
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P

    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError("Parent must be Document or Cell")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield parent.paragraphs[parent.paragraphs.index(child.getparent().paragraphs[0])] # This is wrong logic for parent.paragraphs
        # Actually python-docx has a better way
        pass

# Re-evaluating iter_paragraphs logic for python-docx
def _iter_block_items(parent):
    """
    Duyệt qua các thành phần (Paragraph hoặc Table) theo đúng thứ tự xuất hiện.
    """
    from docx.document import Document as _Document
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph

    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError("Parent must be Document or Cell")

    for child in parent_elm.iterchildren():
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def get_all_paragraphs(parent):
    """Lấy tất cả paragraphs bao gồm cả trong table."""
    paragraphs = []
    for item in _iter_block_items(parent):
        from docx.text.paragraph import Paragraph
        from docx.table import Table
        if isinstance(item, Paragraph):
            paragraphs.append(item)
        elif isinstance(item, Table):
            for row in item.rows:
                for cell in row.cells:
                    paragraphs.extend(get_all_paragraphs(cell))
    return paragraphs

def extract_text(file_path: str) -> list[str]:
    """Trích xuất từng đoạn văn (bao gồm cả trong bảng) dưới dạng Markdown."""
    doc = Document(file_path)
    all_ps = get_all_paragraphs(doc)
    md_paragraphs = []
    for p in all_ps:
        if p.text.strip():
            md = paragraph_to_markdown(p)
            if md:
                md_paragraphs.append(md)
    return md_paragraphs

def reconstruct_document(original_path: str, translated_paragraphs: list[str], output_path: str) -> None:
    """
    Tái tạo tài liệu bằng cách thay thế văn bản đã dịch vào các đoạn văn tương ứng (kể cả trong bảng).
    """
    doc = Document(original_path)
    all_ps = get_all_paragraphs(doc)
    
    non_empty_idx = 0
    for p in all_ps:
        if p.text.strip():
            if non_empty_idx < len(translated_paragraphs):
                translation = translated_paragraphs[non_empty_idx]
                if translation:
                    markdown_to_paragraph(translation, p)
            non_empty_idx += 1
            
    doc.save(output_path)
