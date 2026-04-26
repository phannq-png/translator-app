import customtkinter as ctk
from tkinter import filedialog, messagebox
from core.docx_processor import extract_text, reconstruct_document
from core.chunker import create_fixed_chunks
from core.ai_engine import detect_domain, translate_and_review, extract_terms, get_translation_prompt, get_review_prompt
from core.glossary_manager import get_glossary, save_glossary_batch
from core.config_manager import load_config
from core.progress_manager import save_progress, load_progress
from ui.ui_term_review import TermReviewWindow
import os
import threading
import time
from PIL import Image, ImageTk

try:
    import windnd
except ImportError:
    windnd = None

class MainView(ctk.CTkFrame):
    def __init__(self, master, **kwargs):
        super().__init__(master, **kwargs)
        
        self.file_path = None
        self.original_paragraphs = []
        self.chunks = []
        self.translated_data = {}
        self.domain = ""
        self.current_page = 0
        self.is_loading = False
        self.is_ready = False
        self.spinner_idx = 0
        
        # Load spinner image
        try:
            self.spinner_image_raw = Image.open("assets/spinner.png").convert("RGBA")
        except:
            self.spinner_image_raw = None
        
        # UI: Top bar (Load file & Info)
        self.frame_top = ctk.CTkFrame(self)
        self.frame_top.pack(fill="x", padx=10, pady=10)
        
        self.btn_load = ctk.CTkButton(self.frame_top, text="Chọn file Word (.docx)", command=self.load_file)
        self.btn_load.pack(side="left", padx=10)
        
        self.lbl_file = ctk.CTkLabel(self.frame_top, text="Chưa chọn file")
        self.lbl_file.pack(side="left", padx=10)
        
        self.lbl_domain = ctk.CTkLabel(self.frame_top, text="Lĩnh vực: Đang nhận diện...", text_color=("#D35400", "#E67E22"))
        self.lbl_domain.pack(side="left", padx=20)
        
        # UI: Bottom bar (Action buttons) - Cố định ở dưới cùng
        self.frame_bottom = ctk.CTkFrame(self)
        self.frame_bottom.pack(side="bottom", fill="x", padx=10, pady=10)
        
        self.lbl_start_at = ctk.CTkLabel(self.frame_bottom, text="Đã nạp:")
        self.lbl_start_at.pack(side="left", padx=(10, 2))
        
        self.lbl_total_chunks = ctk.CTkLabel(self.frame_bottom, text="0 khối", text_color=("#1F6AA5", "#3B8ED0"))
        self.lbl_total_chunks.pack(side="left", padx=2)
        
        self.btn_save_draft = ctk.CTkButton(self.frame_bottom, text="Lưu Bản Nháp", command=self.save_draft, state="disabled", fg_color="gray")
        self.btn_save_draft.pack(side="left", padx=10)
        
        self.btn_export = ctk.CTkButton(self.frame_bottom, text="Xác nhận & Xuất File", command=self.export_file, state="disabled", fg_color="green")
        self.btn_export.pack(side="right", padx=10)
        
        self.lbl_status = ctk.CTkLabel(self.frame_bottom, text="Sẵn sàng")
        self.lbl_status.pack(side="left", padx=20, expand=True)

        # Progress Bar - Nằm ngay trên Footer
        self.progress_load = ctk.CTkProgressBar(self)
        self.progress_load.pack(side="bottom", fill="x", padx=20, pady=5)
        self.progress_load.set(0)
        self.progress_load.pack_forget()

        # UI: Bảng đối chiếu (Single Page View) - Chiếm phần còn lại ở giữa
        self.frame_center = ctk.CTkFrame(self)
        self.frame_center.pack(side="top", fill="both", expand=True, padx=10, pady=5)
        
        self.frame_page = ctk.CTkFrame(self.frame_center)
        self.frame_page.pack(side="left", fill="both", expand=True)
        
        # UI: Sidebar Thuật ngữ (Glossary Sidebar)
        self.frame_glossary = ctk.CTkScrollableFrame(self.frame_center, width=250, label_text="Thuật ngữ trong văn bản")
        self.frame_glossary.pack(side="right", fill="y", padx=(5, 0))
        
        # --- LOADING OVERLAY (Dạng quay quay) ---
        self.frame_loading = ctk.CTkFrame(self.frame_page, fg_color=("white", "black"))
        
        self.lbl_spinner = ctk.CTkLabel(self.frame_loading, text="")
        self.lbl_spinner.place(relx=0.5, rely=0.4, anchor="center")
        
        self.lbl_loading_text = ctk.CTkLabel(self.frame_loading, text="Đang xử lý dữ liệu...", font=("Arial", 14))
        self.lbl_loading_text.place(relx=0.5, rely=0.55, anchor="center")
        
        self.progress_inner = ctk.CTkProgressBar(self.frame_loading, width=300)
        self.progress_inner.place(relx=0.5, rely=0.65, anchor="center")
        # ----------------------------------------
        
        # Header trang
        self.lbl_page_info = ctk.CTkLabel(self.frame_page, text="Chưa có dữ liệu", font=("Arial", 14, "bold"), text_color=("#1F6AA5", "#3B8ED0"))
        self.lbl_page_info.pack(side="top", pady=5)
        
        # Điều hướng trang - Ghim ở dưới cùng của frame_page
        self.frame_nav = ctk.CTkFrame(self.frame_page, fg_color="transparent")
        self.frame_nav.pack(side="bottom", fill="x", pady=10)
        
        self.btn_prev = ctk.CTkButton(self.frame_nav, text="< Trang trước", command=self.prev_page, width=100)
        self.btn_prev.pack(side="left", padx=50)
        
        self.lbl_page_num = ctk.CTkLabel(self.frame_nav, text="Trang 0 / 0")
        self.lbl_page_num.pack(side="left", expand=True)
        
        self.btn_translate_page = ctk.CTkButton(self.frame_nav, text="🚀 Dịch trang", command=self.translate_current_page, state="disabled", fg_color="blue", width=100)
        self.btn_translate_page.pack(side="left", padx=5)
        
        self.btn_batch_translate = ctk.CTkButton(self.frame_nav, text="⏩ Dịch tất cả", command=self.start_batch_translation, state="disabled", fg_color="darkblue", width=100)
        self.btn_batch_translate.pack(side="left", padx=5)
        
        self.btn_get_prompt = ctk.CTkButton(self.frame_nav, text="📋 Lấy Prompt", command=self.get_prompt_manual, state="disabled", fg_color="gray", width=100)
        self.btn_get_prompt.pack(side="left", padx=5)
        
        self.btn_review_prompt = ctk.CTkButton(self.frame_nav, text="🔍 Review Prompt", command=self.get_review_prompt_manual, state="disabled", fg_color="purple", width=120)
        self.btn_review_prompt.pack(side="left", padx=5)
        
        self.btn_next = ctk.CTkButton(self.frame_nav, text="Trang sau >", command=self.next_page, width=100)
        self.btn_next.pack(side="right", padx=50)

        # Ô tiếng Nhật (Read-only) - Chiếm phần trên
        self.txt_ja = ctk.CTkTextbox(self.frame_page, height=200, fg_color="transparent", border_width=1)
        self.txt_ja.pack(side="top", fill="both", expand=True, padx=10, pady=5)
        
        # Ô tiếng Việt (Editable) - Chiếm phần dưới
        self.txt_vi = ctk.CTkTextbox(self.frame_page, height=200)
        self.txt_vi.pack(side="top", fill="both", expand=True, padx=10, pady=5)
        
        # Đăng ký kéo thả (nếu có thư viện)
        if windnd:
            windnd.hook_dropfiles(self, self._on_file_drop)

    def _on_file_drop(self, files):
        if not files: return
        file_path = files[0].decode('utf-8') if isinstance(files[0], bytes) else files[0]
        if file_path.lower().endswith(".docx"):
            self.file_path = file_path
            self.lbl_file.configure(text=os.path.basename(self.file_path))
            self.load_file_direct(file_path)
            
    def load_file_direct(self, path):
        self.file_path = path
        self.lbl_status.configure(text="Đang bóc tách văn bản...")
        self.frame_loading.place(relx=0, rely=0, relwidth=1, relheight=1)
        self.frame_loading.lift()
        self.is_loading = True
        self.spinner_idx = 0
        self._animate_spinner()
        threading.Thread(target=self._async_load_process, daemon=True).start()

    def load_file(self):
        self.file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if self.file_path:
            self.lbl_file.configure(text=os.path.basename(self.file_path))
            self.lbl_status.configure(text="Đang bóc tách văn bản...")
            
            # Hiện màn hình loading quay quay
            self.frame_loading.place(relx=0, rely=0, relwidth=1, relheight=1)
            self.frame_loading.lift() # Đưa lên trên cùng của frame_page
            self.is_loading = True
            self.spinner_idx = 0
            self._animate_spinner()
            
            # Chạy nền để tránh treo UI
            threading.Thread(target=self._async_load_process, daemon=True).start()

    def _animate_spinner(self):
        if not self.is_loading: return
        
        if self.spinner_image_raw:
            # Xoay ảnh (mỗi bước 30 độ)
            angle = (self.spinner_idx * 30) % 360
            rotated_img = self.spinner_image_raw.rotate(-angle)
            
            # Cập nhật ảnh cho Label
            spinner_ctk = ctk.CTkImage(light_image=rotated_img, dark_image=rotated_img, size=(60, 60))
            self.lbl_spinner.configure(image=spinner_ctk, text="")
        else:
            # Fallback nếu không load được ảnh: dùng ký hiệu văn bản
            symbols = ["◐", "◓", "◑", "◒"]
            self.lbl_spinner.configure(text=symbols[self.spinner_idx % 4], font=("Arial", 40))
        
        self.spinner_idx += 1
        self.after(80, self._animate_spinner)

    def _async_load_process(self):
        try:
            # 1. Bóc tách Docx sang Markdown (từng đoạn)
            # Vì extract_text gọi pypandoc nhiều lần, ta cần quan sát tiến trình
            from docx import Document
            doc = Document(self.file_path)
            total_p = len(doc.paragraphs)
            
            from core.docx_processor import paragraph_to_markdown
            self.original_paragraphs = []
            for i, p in enumerate(doc.paragraphs):
                if p.text.strip():
                    md = paragraph_to_markdown(p)
                    if md: self.original_paragraphs.append(md)
                
                # Cập nhật progress
                self.progress_inner.set((i + 1) / total_p)
                self.after(0, lambda: self.lbl_loading_text.configure(text=f"Đang bóc tách: {i+1}/{total_p} đoạn..."))
            
            # 2. Gom nhóm khối
            config = load_config()
            batch_size = config.get("batch_size", 5)
            # Chuyển sang gom nhóm theo số đoạn văn cố định
            self.chunks = create_fixed_chunks(self.original_paragraphs, min_sentences=batch_size)
            
            # 3. Tải tiến trình cũ
            progress = load_progress(self.file_path)
            if progress:
                self.domain = progress.get("domain", "")
                self.translated_data = progress.get("translations", {})
                self.after(0, lambda: self.lbl_domain.configure(text=f"Lĩnh vực: {self.domain} (Khôi phục)"))
                self.after(0, lambda: self.lbl_status.configure(text="✅ Đã khôi phục tiến trình cũ."))
            else:
                self.domain = ""
                self.translated_data = {}
            
            # 4. Hiệu ứng hoàn tất mượt mà
            self.after(0, lambda: self.progress_inner.configure(progress_color="green"))
            self.after(0, lambda: self.lbl_loading_text.configure(text="✅ Thành công! Đang chuẩn bị giao diện..."))
            time.sleep(0.5)
            
            # 5. Hiển thị trang đầu tiên
            self.is_loading = False
            self.current_page = 0
            self.after(0, self._show_current_page)
            self.after(200, lambda: self.frame_loading.place_forget())
            self.after(200, lambda: self.lbl_status.configure(text=f"Sẵn sàng: {len(self.chunks)} trang."))
            
            if not self.domain:
                threading.Thread(target=self._async_detect_domain, daemon=True).start()
            else:
                # Nếu đã có domain (khôi phục tiến trình), mở khóa UI luôn
                self.after(0, self.enable_translation_ui)
                
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Lỗi", f"Lỗi nạp file: {e}"))
            self.after(0, lambda: self.progress_load.pack_forget())

    def _show_current_page(self):
        """Hiển thị nội dung của chunk hiện tại lên UI."""
        if not self.chunks: return
        
        chunk = self.chunks[self.current_page]
        chunk_text = chunk['text']
        p_range = chunk['p_range']
        
        # Cập nhật thông tin trang
        range_str = f"Đoạn {p_range[0]+1}" if p_range[0] == p_range[1] else f"Đoạn {p_range[0]+1} - {p_range[1]+1}"
        self.lbl_page_info.configure(text=f"KHỐI VĂN BẢN {self.current_page + 1} ({range_str})", text_color=("#1F6AA5", "#3B8ED0"))
        self.lbl_page_num.configure(text=f"Trang {self.current_page + 1} / {len(self.chunks)}")
        
        # Hiển thị Tiếng Nhật
        self.txt_ja.configure(state="normal")
        self.txt_ja.delete("1.0", "end")
        self.txt_ja.insert("1.0", chunk_text)
        self.txt_ja.configure(state="disabled")
        
        # Hiển thị Tiếng Việt
        self.txt_vi.delete("1.0", "end")
        has_translation = False
        if chunk_text in self.translated_data:
            self.txt_vi.insert("1.0", self.translated_data[chunk_text])
            has_translation = True
        
        # Cập nhật sidebar thuật ngữ tương ứng với trang này
        self._update_glossary_sidebar()
        
        # Trạng thái nút
        self.btn_prev.configure(state="normal" if self.current_page > 0 else "disabled")
        self.btn_next.configure(state="normal" if self.current_page < len(self.chunks)-1 else "disabled")
        
        # Các nút dịch và prompt chỉ hiện khi đã sẵn sàng hoặc đã có bản dịch cũ
        if self.is_ready or has_translation:
            self.btn_translate_page.configure(state="normal")
            self.btn_batch_translate.configure(state="normal")
            self.btn_get_prompt.configure(state="normal")
            self.btn_review_prompt.configure(state="normal" if has_translation else "disabled")
            self.btn_save_draft.configure(state="normal")
            self.btn_export.configure(state="normal")
        else:
            self.btn_translate_page.configure(state="disabled")
            self.btn_batch_translate.configure(state="disabled")
            self.btn_get_prompt.configure(state="disabled")
            self.btn_review_prompt.configure(state="disabled")
            self.btn_save_draft.configure(state="disabled")
            self.btn_export.configure(state="disabled")

    def next_page(self):
        # Lưu dữ liệu trang hiện tại
        self._save_current_page_to_memory()
        if self.current_page < len(self.chunks) - 1:
            self.current_page += 1
            self._show_current_page()

    def prev_page(self):
        self._save_current_page_to_memory()
        if self.current_page > 0:
            self.current_page -= 1
            self._show_current_page()

    def _save_current_page_to_memory(self):
        """Lưu text từ ô Tiếng Việt vào bộ nhớ tạm."""
        if not self.chunks: return
        chunk_text = self.chunks[self.current_page]['text']
        translated_text = self.txt_vi.get("1.0", "end-1c").strip()
        if translated_text:
            self.translated_data[chunk_text] = translated_text

    def _async_detect_domain(self):
        try:
            # Lấy khối văn bản đầu tiên để nhận diện (tối đa 1000 ký tự)
            if self.chunks:
                sample_text = self.chunks[0]['text'][:1000]
                self.domain = detect_domain(sample_text)
                self.after(0, lambda: self.lbl_domain.configure(text=f"Lĩnh vực: {self.domain}", text_color=("#27AE60", "#2ECC71")))
                self.after(0, lambda: self.lbl_status.configure(text=f"Đã nhận diện: {self.domain}"))
                # Tự động trích xuất thuật ngữ ngay sau khi có lĩnh vực
                self.after(500, self.start_extraction)
        except Exception as e:
            self.after(0, lambda: self.lbl_status.configure(text=f"Lỗi API nhận diện: {str(e)[:80]}..."))

    def start_extraction(self):
        if not self.chunks: return
        self.lbl_status.configure(text="Đang tự động phân tích và trích xuất thuật ngữ chuyên ngành...")
        threading.Thread(target=self._async_extract_terms, daemon=True).start()

    def _async_extract_terms(self):
        try:
            # Lấy khoảng 2000 ký tự đầu tiên để trích xuất
            sample_text = ""
            for chunk in self.chunks[:3]: 
                sample_text += chunk['text'] + "\n"
            
            terms = extract_terms(sample_text, self.domain)
            
            if terms:
                from core.glossary_manager import save_glossary_batch
                save_glossary_batch(self.domain, terms)
                self.after(0, lambda: self.lbl_status.configure(text=f"Đã trích xuất {len(terms)} thuật ngữ. Đang chờ review..."))
            else:
                self.after(0, lambda: self.lbl_status.configure(text="Không tìm thấy thuật ngữ mới. Đang chờ review..."))
            
            # Hiển thị màn hình Review
            self.after(500, self.show_term_review)
                
        except Exception as e:
            self.after(0, lambda: self.lbl_status.configure(text=f"Lỗi trích xuất AI: {str(e)[:50]}"))
        finally:
            pass

    def show_term_review(self):
        """Hiển thị màn hình review thuật ngữ."""
        review_win = TermReviewWindow(
            self, 
            self.domain, 
            self.chunks, 
            start_translation_callback=self.enable_translation_ui
        )
        # Sidebar cũng cần cập nhật
        self._update_glossary_sidebar()

    def enable_translation_ui(self):
        """Mở khóa các nút dịch sau khi đã review thuật ngữ."""
        self.is_ready = True
        self.btn_translate_page.configure(state="normal")
        self.btn_batch_translate.configure(state="normal")
        self.btn_get_prompt.configure(state="normal")
        self.btn_export.configure(state="normal")
        self.btn_save_draft.configure(state="normal")
        self.lbl_status.configure(text="Sẵn sàng. Bạn có thể dịch từng trang hoặc toàn bộ.")
        # Cập nhật sidebar ngay khi mở khóa
        self.after(0, self._update_glossary_sidebar)

    def translate_current_page(self):
        if not self.chunks: return
        self.btn_translate_page.configure(state="disabled", text="⌛ Đang dịch...")
        threading.Thread(target=self._async_translate_page, daemon=True).start()

    def _async_translate_page(self):
        try:
            chunk = self.chunks[self.current_page]
            glossary = get_glossary(self.domain)
            results = translate_and_review([chunk['text']], self.domain, glossary)
            
            if results:
                self.translated_data[chunk['text']] = results[0]
                self.after(0, self._show_current_page)
                self.after(0, lambda: self.lbl_status.configure(text=f"Đã dịch xong trang {self.current_page+1}"))
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Lỗi dịch trang", str(e)))
        finally:
            self.after(0, lambda: self.btn_translate_page.configure(state="normal", text="🚀 Dịch trang"))

    def get_prompt_manual(self):
        if not self.chunks: return
        chunk = self.chunks[self.current_page]
        glossary = get_glossary(self.domain)
        prompt = get_translation_prompt([chunk['text']], self.domain, glossary)
        self._show_prompt_popup("Copy Dịch Prompt", prompt)

    def get_review_prompt_manual(self):
        if not self.chunks: return
        chunk = self.chunks[self.current_page]
        translated_text = self.txt_vi.get("1.0", "end-1c").strip()
        if not translated_text:
            messagebox.showwarning("Thông báo", "Vui lòng dịch trang này trước khi lấy Prompt Review.")
            return
            
        glossary = get_glossary(self.domain)
        prompt = get_review_prompt([chunk['text']], [translated_text], self.domain, glossary)
        self._show_prompt_popup("Copy Review Prompt", prompt)

    def _show_prompt_popup(self, title, content):
        pop = ctk.CTkToplevel(self)
        pop.title(title)
        pop.geometry("600x500")
        pop.attributes("-topmost", True)
        
        txt = ctk.CTkTextbox(pop)
        txt.pack(fill="both", expand=True, padx=10, pady=10)
        txt.insert("1.0", content)
        
        btn_copy = ctk.CTkButton(pop, text="Copy vào Clipboard", 
                                 command=lambda: self._copy_to_clipboard(content, pop))
        btn_copy.pack(pady=10)

    def _copy_to_clipboard(self, text, window):
        self.clipboard_clear()
        self.clipboard_append(text)
        messagebox.showinfo("OK", "Đã copy vào clipboard!", parent=window)

    def _update_glossary_sidebar(self):
        """Cập nhật danh sách thuật ngữ ở sidebar theo trang hiện tại."""
        if not hasattr(self, 'frame_glossary'): return
        
        # Xóa các widget cũ
        for child in self.frame_glossary.winfo_children():
            child.destroy()
            
        if not self.chunks: return
        
        # Lấy văn bản của trang hiện tại
        current_text = self.chunks[self.current_page]['text']
        
        active_terms = {}
        if self.domain:
            from core.glossary_manager import get_glossary
            all_glossary = get_glossary(self.domain)
            # Chỉ hiển thị những từ xuất hiện trong văn bản hiện tại
            for ja, vi in all_glossary.items():
                ja_clean = ja.strip()
                if ja_clean in current_text:
                    active_terms[ja_clean] = vi.strip()
        
        if not active_terms:
            lbl = ctk.CTkLabel(self.frame_glossary, text="(Không có thuật ngữ)", text_color="gray", font=("Arial", 11, "italic"))
            lbl.pack(pady=20)
            return

        for ja, vi in active_terms.items():
            row = ctk.CTkFrame(self.frame_glossary, fg_color="transparent")
            row.pack(fill="x", pady=5)
            
            lbl_ja = ctk.CTkLabel(row, text=ja, font=("Arial", 11, "bold"), anchor="w")
            lbl_ja.pack(fill="x")
            
            ent_vi = ctk.CTkEntry(row, height=25, font=("Arial", 11))
            ent_vi.insert(0, vi)
            ent_vi.pack(side="left", fill="x", expand=True)
            
            btn_save = ctk.CTkButton(row, text="ok", width=30, height=25, 
                                     command=lambda j=ja, e=ent_vi: self._quick_save_term(j, e.get()))
            btn_save.pack(side="right", padx=2)

    def _quick_save_term(self, ja, vi):
        from core.glossary_manager import update_term
        update_term(ja, vi, self.domain)
        self.lbl_status.configure(text=f"Đã cập nhật thuật ngữ: {ja}")

    def start_batch_translation(self):
        """Bắt đầu dịch hàng loạt tất cả các trang chưa dịch."""
        if not self.chunks:
            return
            
        if not messagebox.askyesno("Xác nhận", "Bạn có muốn bắt đầu dịch hàng loạt tất cả các trang chưa được dịch?"):
            return
            
        self.btn_batch_translate.configure(state="disabled", text="⌛ Đang dịch...")
        self.btn_translate_page.configure(state="disabled")
        self.lbl_status.configure(text="Đang dịch tự động toàn bộ văn bản...")
        
        threading.Thread(target=self._async_translation_loop, daemon=True).start()
        
    def _async_translation_loop(self):
        """Luồng dịch tự động chạy qua từng trang."""
        glossary = get_glossary(self.domain)
        
        # Bắt đầu dịch từ trang hiện tại hoặc trang đầu tiên chưa dịch
        for i in range(len(self.chunks)):
            chunk = self.chunks[i]
            chunk_text = chunk['text']
            
            # Kiểm tra nếu trang này đã dịch rồi thì bỏ qua
            if chunk_text in self.translated_data and self.translated_data[chunk_text].strip():
                continue
            
            # Chuyển trang UI để người dùng thấy tiến trình (nếu muốn)
            self.current_page = i
            self.after(0, self._show_current_page)
            
            # Hiển thị trạng thái đang dịch trên ô Tiếng Việt của trang hiện tại
            self.after(0, lambda idx=i: self.lbl_status.configure(text=f"Đang dịch Trang {idx+1}/{len(self.chunks)}..."))
            
            success = False
            retry_count = 0
            while not success and retry_count < 3:
                try:
                    results = translate_and_review([chunk_text], self.domain, glossary)
                    final_vi = results[0]
                    
                    self.translated_data[chunk_text] = final_vi
                    self.after(0, self._show_current_page) # Cập nhật lại UI trang đó
                    
                    # Tự động cuộn xuống cuối ô tiếng Việt để thấy kết quả mới nhất
                    self.after(50, lambda: self.txt_vi.see("end"))
                    
                    save_progress(self.file_path, self.domain, self.translated_data)
                    if hasattr(self.master, 'update_recent_sidebar'):
                        self.after(0, self.master.update_recent_sidebar)
                    success = True
                except Exception as e:
                    retry_count += 1
                    if retry_count >= 3:
                        res = messagebox.askyesno("Lỗi API", f"Lỗi liên tiếp tại trang {i+1}:\n{e}\n\nBạn có muốn thử lại lần nữa không?")
                        if not res:
                            self.after(0, self._finish_batch_translation)
                            return
                        retry_count = 0 # Reset để thử lại nếu người dùng đồng ý
                    time.sleep(2) # Đợi một chút trước khi thử lại
            
            # Nghỉ một chút giữa các request để tránh rate limit
            time.sleep(1)
                
        self.after(0, self._finish_batch_translation)

    def _finish_batch_translation(self):
        """Kết thúc quá trình dịch hàng loạt."""
        self.lbl_status.configure(text="Hoàn tất dịch hàng loạt.")
        self.btn_batch_translate.configure(state="normal", text="⏩ Dịch tất cả")
        self.btn_translate_page.configure(state="normal")
        messagebox.showinfo("Hoàn tất", "Đã hoàn thành dịch hàng loạt các trang chưa dịch!")

    def save_draft(self):
        if not self.file_path: return
        self._save_current_page_to_memory() # Lưu trang đang đứng trước khi ghi file
        save_progress(self.file_path, self.domain, self.translated_data)
        if hasattr(self.master, 'update_recent_sidebar'):
            self.master.update_recent_sidebar()
        messagebox.showinfo("Thông báo", "Đã lưu bản nháp!")

    def _get_reassembled_data(self):
        """Lắp ráp lại các đoạn văn gốc từ các khối đã dịch."""
        self._save_current_page_to_memory() # Đảm bảo trang hiện tại được lưu vào memory
        
        para_map = {}
        for i, chunk in enumerate(self.chunks):
            p_range = chunk['p_range']
            trans = self.translated_data.get(chunk['text'], "").strip()
            
            if p_range[0] == p_range[1]:
                p_idx = p_range[0]
                if p_idx not in para_map: para_map[p_idx] = trans
                else: para_map[p_idx] += " " + trans
            else:
                # Nếu chunk gồm nhiều đoạn văn, ta tách ra theo dấu \n\n
                sub_trans = trans.split("\n\n")
                for offset, p_idx in enumerate(range(p_range[0], p_range[1] + 1)):
                    para_map[p_idx] = sub_trans[offset] if offset < len(sub_trans) else ""
        
        para_list = []
        for idx in range(len(self.original_paragraphs)):
            para_list.append(para_map.get(idx, ""))
            
        return para_list

    def export_file(self):
        if not self.file_path or not self.original_paragraphs: return
        
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"{os.path.splitext(os.path.basename(self.file_path))[0]}_translated.docx",
            filetypes=[("Word Documents", "*.docx")]
        )
        
        if save_path:
            try:
                self.lbl_status.configure(text="Đang lắp ráp và xuất file...")
                translated_paragraphs = self._get_reassembled_data()
                reconstruct_document(self.file_path, translated_paragraphs, save_path)
                self.lbl_status.configure(text="Lưu file thành công!")
                messagebox.showinfo("Thành công", "Đã xuất file dịch!")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Lỗi: {e}")
