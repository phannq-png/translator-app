import os
import pytest
from unittest.mock import patch, MagicMock
from docx import Document
from core.docx_processor import extract_text, reconstruct_document

# Mocking pypandoc because it might not be installed in the environment
# and it requires external pandoc binary.
@pytest.fixture
def mock_pypandoc():
    with patch('core.docx_processor.pypandoc') as mock:
        # Mock convert_file to return some markdown for docx
        mock.convert_file.side_effect = lambda path, to, format=None, outputfile=None: "Mocked Markdown content"
        yield mock

def test_extract_text_basic(tmp_path, mock_pypandoc):
    # Create a dummy docx file
    test_docx = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Hello World")
    doc.save(test_docx)
    
    with patch('core.docx_processor.paragraph_to_markdown') as mock_p2m:
        mock_p2m.return_value = "Hello World MD"
        texts = extract_text(str(test_docx))
        
        assert len(texts) == 1
        assert texts[0] == "Hello World MD"

def test_reconstruct_document_with_mock(tmp_path, mock_pypandoc):
    # Create a dummy docx file with specific formatting
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("Original Text")
    r.bold = True
    doc.save(input_docx)
    
    translated_texts = ["Bản dịch tiếng Việt"]
    
    # Mock markdown_to_paragraph to avoid real pandoc calls
    with patch('core.docx_processor.markdown_to_paragraph') as mock_m2p:
        reconstruct_document(str(input_docx), translated_texts, str(output_docx))
        
        assert mock_m2p.called
        # Check if output file exists
        assert os.path.exists(output_docx)

def test_integration_with_real_file_structure():
    """
    Kịch bản test sử dụng file thực tế trong folder test/input
    Lưu ý: Chỉ test được phần extract_text nếu không có pandoc.
    """
    input_file = "test/input/[JP]BMT 54-2026.docx"
    if not os.path.exists(input_file):
        pytest.skip("File test/input/[JP]BMT 54-2026.docx không tồn tại")
        
    try:
        # Test trích xuất text (Mocking pandoc)
        with patch('core.docx_processor.pypandoc.convert_file') as mock_conv:
            mock_conv.return_value = "Mocked Markdown"
            texts = extract_text(input_file)
            print(f"Extracted {len(texts)} paragraphs from real file.")
            assert len(texts) > 0
    except Exception as e:
        pytest.fail(f"Lỗi khi xử lý file thực tế: {e}")

if __name__ == "__main__":
    # Cho phép chạy file này trực tiếp để test nhanh
    pytest.main([__file__])
